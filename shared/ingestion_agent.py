"""Document ingestion agent for SHIMS Enterprise.

Parses Word, PDF, Excel, and text files to extract structured experiment data,
SOP content, or BMR information. Uses LLM for structured extraction when
simple text parsing is insufficient.
"""
from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from .config import settings
from .database import db

try:
    from .ai import ask_ai, extract_json_maybe
    _AI_OK = True
except Exception:
    _AI_OK = False
    ask_ai = None  # type: ignore


def _clean(text: Any, limit: int = 4000) -> str:
    return re.sub(r'\s+', ' ', str(text or '')).strip()[:limit]


def _load_json(text: str | None, default: Any) -> Any:
    if not text:
        return default
    try:
        return json.loads(text)
    except Exception:
        return default


def extract_text_from_file(path: str | Path) -> str:
    """Extract raw text from Word, PDF, Excel, or text files."""
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix in {'.docx', '.doc'}:
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(str(p))
            parts = ['\n'.join(p.text for p in doc.paragraphs if p.text.strip())]
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    vals = [_clean(cell.text, 500) for cell in row.cells]
                    if any(vals):
                        rows.append(' | '.join(vals))
                if rows:
                    parts.append('\n'.join(rows))
            return '\n\n'.join(parts)
        except Exception as exc:
            return f'[docx extraction failed: {exc}]'
    if suffix == '.pdf':
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(p))
            return '\n\n'.join(page.extract_text() or '' for page in reader.pages)
        except Exception as exc:
            return f'[pdf extraction failed: {exc}]'
    if suffix in {'.xlsx', '.xls'}:
        try:
            from openpyxl import load_workbook
            wb = load_workbook(str(p), data_only=True, read_only=True)
            parts = []
            for sheet in wb:
                rows = []
                for row in sheet.iter_rows(values_only=True):
                    vals = [str(v) if v is not None else '' for v in row]
                    if any(vals):
                        rows.append(' | '.join(vals))
                if rows:
                    parts.append(f'[sheet:{sheet.title}]\n' + '\n'.join(rows))
            return '\n\n'.join(parts)
        except Exception as exc:
            return f'[xlsx extraction failed: {exc}]'
    if suffix in {'.txt', '.md', '.csv'}:
        return p.read_text(encoding='utf-8', errors='ignore')
    return f'[unsupported format: {suffix}]'


EXTRACTION_SYSTEM = """You are a pharmaceutical document parsing AI. Extract structured experiment data from the provided text.

Return ONLY a JSON object with this exact structure:
{
  "product_name": "...",
  "document_type": "experiment_report|sop|bmr|flow_chart|other",
  "stages": [
    {
      "stage_no": 1,
      "stage_name": "reaction|quench|extraction|wash|crystallization|centrifugation|filtration|drying|sizing|milling|blending|packing|packaging|labeling",
      "description": "brief description",
      "temperature_c": null,
      "ph_value": null,
      "solvent": "",
      "catalyst": "",
      "rm_description": "raw materials used",
      "theoretical_yield_pct": null,
      "actual_yield_pct": null,
      "input_qty": null,
      "output_qty": null,
      "purity_pct": null,
      "notes": ""
    }
  ],
  "summary": {
    "batch_size": "",
    "overall_yield": "",
    "key_observations": "",
    "qc_tests": []
  }
}

Rules:
- stage_name MUST be one of the canonical names listed above.
- Use null for numeric fields that are not present in the text.
- Do not invent data not present in the source.
- If the text is an SOP, return empty stages and fill summary with SOP metadata.
"""


def parse_document_to_experiment(path: str | Path) -> dict[str, Any]:
    """Parse a document into structured experiment data."""
    text = extract_text_from_file(path)
    if not text or len(text) < 30:
        return {'ok': False, 'error': 'Could not extract meaningful text from file'}

    # Try simple keyword extraction first
    product = _detect_product_name(text)
    stages = _extract_stages_heuristic(text)

    # If we have good heuristic data, return it
    if stages and len(stages) >= 3:
        return {
            'ok': True,
            'source': 'heuristic',
            'product_name': product or 'Unknown',
            'stages': stages,
            'raw_text_preview': text[:2000],
        }

    # Fallback to LLM extraction
    if not _AI_OK or ask_ai is None:
        return {
            'ok': True,
            'source': 'heuristic_partial',
            'product_name': product or 'Unknown',
            'stages': stages or [],
            'raw_text_preview': text[:2000],
            'note': 'LLM unavailable for deep extraction',
        }

    import asyncio
    prompt = f"Extract structured experiment data from the following pharmaceutical document:\n\n{text[:15000]}\n\nReturn JSON only."
    try:
        result = asyncio.run(ask_ai(prompt, system=EXTRACTION_SYSTEM, provider=settings.ai_provider, model=settings.ollama_model))
        parsed = extract_json_maybe(result.text)
        if isinstance(parsed, dict) and parsed.get('stages'):
            return {
                'ok': True,
                'source': 'llm',
                'product_name': parsed.get('product_name', product or 'Unknown'),
                'stages': parsed.get('stages', []),
                'summary': parsed.get('summary', {}),
                'raw_text_preview': text[:1000],
            }
    except Exception as exc:
        pass

    return {
        'ok': True,
        'source': 'heuristic_fallback',
        'product_name': product or 'Unknown',
        'stages': stages or [],
        'raw_text_preview': text[:2000],
    }


def _detect_product_name(text: str) -> str | None:
    """Try to detect product name from common patterns."""
    patterns = [
        r'(?i)\b(product name|product|api|substance)\s*[:\-]?\s*([A-Za-z][A-Za-z0-9\s\-]{2,40})',
        r'(?i)\b(title of experiment|experiment title)\s*[:\-]?\s*([A-Za-z][A-Za-z0-9\s\-]{2,40})',
        r'(?i)\b(bmr|batch manufacturing record)\s*[:\-]?\s*for\s+([A-Za-z][A-Za-z0-9\s\-]{2,40})',
    ]
    for pat in patterns:
        m = re.search(pat, text)
        if m:
            return _clean(m.group(2))
    # Try first capitalized line
    for line in text.splitlines()[:10]:
        line = line.strip()
        if len(line) > 3 and len(line) < 60 and line[0].isupper() and line.replace(' ', '').isalpha():
            return line
    return None


def _extract_stages_heuristic(text: str) -> list[dict[str, Any]]:
    """Extract stages using keyword heuristics."""
    stage_keywords = {
        'reaction': ['reaction', 'synthesis', 'condensation', 'cyclization', 'alkylation'],
        'quench': ['quench', 'cooling', 'ice bath'],
        'extraction': ['extraction', 'workup', 'separation', 'layer separation'],
        'wash': ['wash', 'washing', 'water wash', 'brine wash'],
        'crystallization': ['crystallization', 'recrystallization', 'seed', 'cooling crystallization'],
        'centrifugation': ['centrifuge', 'centrifugation', 'spin drying'],
        'filtration': ['filtration', 'filter', 'nutsche', 'candle filter', 'leaf filter'],
        'drying': ['drying', 'vacuum dry', 'tray dry', 'fluid bed drying'],
        'sizing': ['sizing', 'sieve', 'mesh'],
        'milling': ['milling', 'mill', 'micronization'],
        'blending': ['blending', 'blend', 'mixer'],
        'packing': ['packing', 'filling', 'drum'],
        'packaging': ['packaging', 'bottle', 'sachet', 'strip'],
        'labeling': ['labeling', 'label', 'coding'],
    }

    lines = text.splitlines()
    stages = []
    stage_no = 0
    seen = set()

    for line in lines:
        line_clean = line.lower()
        for stage_name, keywords in stage_keywords.items():
            if stage_name in seen:
                continue
            for kw in keywords:
                if kw in line_clean:
                    stage_no += 1
                    seen.add(stage_name)
                    # Try to extract temperature
                    temp = None
                    temp_match = re.search(r'(\d+\.?\d*)\s*°?\s*c', line, re.IGNORECASE)
                    if temp_match:
                        temp = float(temp_match.group(1))
                    # Try pH
                    ph = None
                    ph_match = re.search(r'ph\s*[:\-]?\s*(\d+\.?\d*)', line, re.IGNORECASE)
                    if ph_match:
                        ph = float(ph_match.group(1))
                    stages.append({
                        'stage_no': stage_no,
                        'stage_name': stage_name,
                        'description': _clean(line, 300),
                        'temperature_c': temp,
                        'ph_value': ph,
                        'solvent': '',
                        'catalyst': '',
                        'rm_description': '',
                        'theoretical_yield_pct': None,
                        'actual_yield_pct': None,
                        'input_qty': None,
                        'output_qty': None,
                        'purity_pct': None,
                        'notes': _clean(line, 300),
                    })
                    break
        if stage_no >= 15:
            break

    return stages


def ingest_document_as_experiment(file_path: str | Path, user_id: int | None = None) -> dict[str, Any]:
    """Full pipeline: parse document → create experiment → log stages."""
    parsed = parse_document_to_experiment(file_path)
    if not parsed.get('ok'):
        return parsed

    product = parsed.get('product_name') or 'Unknown Product'
    stages = parsed.get('stages', [])

    # Create experiment
    exp_id = db.execute(
        'INSERT INTO rd_experiments(product_name, route_name, notes, status, created_by) VALUES (?, ?, ?, ?, ?)',
        (product, f'{product} Route', f'Auto-ingested from {Path(file_path).name}', 'active', user_id),
    )

    # Log stages
    for st in stages:
        db.execute(
            '''INSERT INTO rd_experiment_stages(
                experiment_id, stage_no, stage_name, temperature_c, ph_value, solvent,
                catalyst, rm_description, theoretical_yield_pct, actual_yield_pct,
                input_qty, output_qty, purity_pct, notes
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
            (
                exp_id,
                st.get('stage_no', 1),
                st.get('stage_name', 'operation'),
                st.get('temperature_c'),
                st.get('ph_value'),
                _clean(st.get('solvent')),
                _clean(st.get('catalyst')),
                _clean(st.get('rm_description')),
                st.get('theoretical_yield_pct'),
                st.get('actual_yield_pct'),
                st.get('input_qty'),
                st.get('output_qty'),
                st.get('purity_pct'),
                _clean(st.get('notes')),
            ),
        )

    # Also ingest into BMR corpus
    from .enterprise_bmr_corpus import import_bmr_folder
    try:
        corpus_result = import_bmr_folder(str(Path(file_path).parent), user_id=user_id, limit=10)
    except Exception:
        corpus_result = {}

    return {
        'ok': True,
        'experiment_id': exp_id,
        'product_name': product,
        'stages_logged': len(stages),
        'source': parsed.get('source'),
        'corpus_imported': corpus_result.get('imported', 0),
    }
