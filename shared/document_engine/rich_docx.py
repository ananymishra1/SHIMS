"""Rich, fully-styleable Word (.docx) document engine for SHIMS.

The earlier `create_docx` could only emit a title + level-2 headings + plain
paragraphs. This engine gives the document creator real editorial control:

  * Named **style profiles** (regulatory / corporate / modern / minimal) that set
    the font family, sizes, colors and spacing for every semantic element —
    title, heading 1/2/3, sub-heading, paragraph, note/callout, list and caption.
  * A **block-based content model** so callers describe structure, not raw runs:
    heading, paragraph, bullets, numbered, note, table, quote, pagebreak, spacer.
  * Per-run / per-paragraph overrides: bold, italic, underline, size, color,
    alignment and left-indent — so numbering, indenting, bullet points, bold and
    size are all individually editable, exactly as requested.
  * Optional **branded letterhead** (logo + company block) and page-numbered
    footer, so generated documents look like company masterpieces.

python-docx is required; if unavailable the caller's higher-level service falls
back to Markdown.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

from .branded_base import COMPANY, _logo_path

_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


@dataclass
class TextStyle:
    """Formatting for one semantic element."""
    font: str = "Calibri"
    size: float = 11
    bold: bool = False
    italic: bool = False
    color: str = "000000"          # hex, no '#'
    align: str = "left"
    space_before: float = 0        # pt
    space_after: float = 6         # pt
    left_indent: float = 0         # inches
    line_spacing: float = 1.15
    all_caps: bool = False
    keep_with_next: bool = False


@dataclass
class DocStyleProfile:
    """A complete, named set of styles for a document family."""
    name: str
    title: TextStyle
    heading1: TextStyle
    heading2: TextStyle
    heading3: TextStyle
    subheading: TextStyle
    paragraph: TextStyle
    note: TextStyle
    list_item: TextStyle
    caption: TextStyle
    note_bg: str = "FEF3C7"        # callout fill
    accent: str = "1E3A8A"
    page_numbers: bool = True


def _ink(c: str = "0F172A") -> TextStyle:  # convenience
    return TextStyle(color=c)


PROFILES: dict[str, DocStyleProfile] = {
    "regulatory": DocStyleProfile(
        name="regulatory",
        title=TextStyle("Times New Roman", 18, bold=True, color="111827", align="center", space_after=10, all_caps=True),
        heading1=TextStyle("Times New Roman", 14, bold=True, color="1E3A8A", space_before=10, space_after=4, keep_with_next=True),
        heading2=TextStyle("Times New Roman", 12, bold=True, color="1E40AF", space_before=8, space_after=3, keep_with_next=True),
        heading3=TextStyle("Times New Roman", 11, bold=True, italic=True, color="334155", space_before=6, space_after=2, keep_with_next=True),
        subheading=TextStyle("Times New Roman", 11, bold=True, color="475569", space_after=3),
        paragraph=TextStyle("Times New Roman", 11, color="111827", align="justify", space_after=6, line_spacing=1.3),
        note=TextStyle("Times New Roman", 10, italic=True, color="92400E", space_after=6),
        list_item=TextStyle("Times New Roman", 11, color="111827", space_after=3),
        caption=TextStyle("Times New Roman", 9, italic=True, color="64748B", align="center", space_after=8),
        accent="1E3A8A",
    ),
    "corporate": DocStyleProfile(
        name="corporate",
        title=TextStyle("Calibri", 22, bold=True, color="0F172A", space_after=12),
        heading1=TextStyle("Calibri", 16, bold=True, color="1D4ED8", space_before=12, space_after=4, keep_with_next=True),
        heading2=TextStyle("Calibri", 13, bold=True, color="2563EB", space_before=8, space_after=3, keep_with_next=True),
        heading3=TextStyle("Calibri", 11.5, bold=True, color="334155", space_before=6, space_after=2, keep_with_next=True),
        subheading=TextStyle("Calibri", 11, bold=True, italic=True, color="475569", space_after=3),
        paragraph=TextStyle("Calibri", 11, color="111827", space_after=8, line_spacing=1.2),
        note=TextStyle("Calibri", 10, color="1E3A8A", space_after=8),
        list_item=TextStyle("Calibri", 11, color="111827", space_after=3),
        caption=TextStyle("Calibri", 9, italic=True, color="64748B", space_after=8),
        note_bg="DBEAFE", accent="2563EB",
    ),
    "modern": DocStyleProfile(
        name="modern",
        title=TextStyle("Segoe UI", 24, bold=True, color="0F766E", space_after=12),
        heading1=TextStyle("Segoe UI Semibold", 16, bold=True, color="0F766E", space_before=12, space_after=4, keep_with_next=True),
        heading2=TextStyle("Segoe UI Semibold", 13, bold=True, color="0D9488", space_before=8, space_after=3, keep_with_next=True),
        heading3=TextStyle("Segoe UI", 11.5, bold=True, color="334155", space_before=6, space_after=2, keep_with_next=True),
        subheading=TextStyle("Segoe UI", 11, italic=True, color="475569", space_after=3),
        paragraph=TextStyle("Segoe UI", 11, color="1F2937", space_after=8, line_spacing=1.25),
        note=TextStyle("Segoe UI", 10, color="115E59", space_after=8),
        list_item=TextStyle("Segoe UI", 11, color="1F2937", space_after=3),
        caption=TextStyle("Segoe UI", 9, italic=True, color="64748B", space_after=8),
        note_bg="CCFBF1", accent="0D9488",
    ),
    "minimal": DocStyleProfile(
        name="minimal",
        title=TextStyle("Calibri", 20, bold=True, color="000000", space_after=10),
        heading1=TextStyle("Calibri", 15, bold=True, color="000000", space_before=10, space_after=4, keep_with_next=True),
        heading2=TextStyle("Calibri", 12.5, bold=True, color="000000", space_before=6, space_after=2, keep_with_next=True),
        heading3=TextStyle("Calibri", 11, bold=True, color="000000", space_before=4, space_after=2, keep_with_next=True),
        subheading=TextStyle("Calibri", 11, italic=True, color="333333", space_after=2),
        paragraph=TextStyle("Calibri", 11, color="000000", space_after=6, line_spacing=1.15),
        note=TextStyle("Calibri", 10, italic=True, color="444444", space_after=6),
        list_item=TextStyle("Calibri", 11, color="000000", space_after=2),
        caption=TextStyle("Calibri", 9, italic=True, color="666666", space_after=6),
        note_bg="F1F5F9", accent="000000",
    ),
}


def _apply(par, style: TextStyle) -> None:
    pf = par.paragraph_format
    pf.space_before = Pt(style.space_before)
    pf.space_after = Pt(style.space_after)
    pf.left_indent = Inches(style.left_indent)
    pf.line_spacing = style.line_spacing
    pf.keep_with_next = style.keep_with_next
    par.alignment = _ALIGN.get(style.align, WD_ALIGN_PARAGRAPH.LEFT)


def _run(par, text: str, base: TextStyle, **ov) -> None:
    r = par.add_run(text)
    f = r.font
    f.name = ov.get("font", base.font)
    f.size = Pt(ov.get("size", base.size))
    f.bold = ov.get("bold", base.bold)
    f.italic = ov.get("italic", base.italic)
    f.underline = ov.get("underline", False)
    f.all_caps = ov.get("all_caps", base.all_caps)
    color = ov.get("color", base.color).lstrip("#")
    f.color.rgb = RGBColor.from_string(color)


def _shade(par, hex_fill: str) -> None:
    """Apply a background fill (callout box) to a paragraph."""
    ppr = par._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill.lstrip("#"))
    ppr.append(shd)


def _add_page_field(par) -> None:
    """Insert a live PAGE field (Word renders the current page number)."""
    run = par.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_end)


def _letterhead(doc, profile: DocStyleProfile, company: dict) -> None:
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo = _logo_path()
    if logo and logo.exists():
        try:
            p.add_run().add_picture(str(logo), height=Inches(0.5))
            p.add_run().add_break()
        except Exception:
            pass
    _run(p, company.get("legal_name", ""), TextStyle(profile.title.font, 12, bold=True, color=profile.accent))
    sub = header.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bits = [company.get("address", "")]
    line2 = "   |   ".join(b for b in [
        f"GSTIN: {company['gstin']}" if company.get("gstin") else "",
        company.get("email", ""), company.get("phone", ""),
    ] if b)
    _run(sub, company.get("address", ""), TextStyle(profile.caption.font, 8.5, color="64748B"))
    if line2:
        sub.add_run().add_break()
        _run(sub, line2, TextStyle(profile.caption.font, 8.5, color="64748B"))

    if profile.page_numbers:
        foot = section.footer.paragraphs[0]
        foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(foot, "Page ", profile.caption)
        _add_page_field(foot)


def _heading_style(profile: DocStyleProfile, level: int) -> TextStyle:
    return {1: profile.heading1, 2: profile.heading2, 3: profile.heading3}.get(level, profile.heading3)


def build_docx(
    title: str,
    blocks: list[dict[str, Any]],
    *,
    profile: str | DocStyleProfile = "corporate",
    output_path: Path | str,
    letterhead: bool = True,
    company: dict | None = None,
    subtitle: str = "",
) -> Path:
    """Render `blocks` into a styled .docx and return the path.

    Each block is a dict with a ``type``:
      - heading:   {type, level(1-3), text}
      - subheading:{type, text}
      - paragraph: {type, text|runs, bold, italic, size, color, align, indent}
      - bullets:   {type, items:[str|{text,level,bold,...}]}
      - numbered:  {type, items:[...]}  (real Word auto-numbering)
      - note:      {type, text}         (shaded callout box)
      - quote:     {type, text}
      - table:     {type, headers:[...], rows:[[...]], widths:[in,...]}
      - caption:   {type, text}
      - pagebreak / spacer
    A paragraph may carry ``runs``: a list of {text,bold,italic,underline,size,color}
    for mixed inline formatting within one line.
    """
    prof = PROFILES.get(profile) if isinstance(profile, str) else profile
    if prof is None:
        prof = PROFILES["corporate"]
    company = company or dict(COMPANY)

    doc = Document()
    if letterhead:
        _letterhead(doc, prof, company)

    # Title + optional subtitle
    tp = doc.add_paragraph()
    _apply(tp, prof.title)
    _run(tp, title, prof.title)
    if subtitle:
        sp = doc.add_paragraph()
        _apply(sp, prof.subheading)
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(sp, subtitle, prof.subheading)

    for blk in blocks:
        bt = (blk.get("type") or "paragraph").lower()

        if bt == "pagebreak":
            doc.add_page_break()
        elif bt == "spacer":
            doc.add_paragraph()
        elif bt == "heading":
            lvl = int(blk.get("level", 1))
            style = _heading_style(prof, lvl)
            p = doc.add_paragraph(); _apply(p, style)
            _run(p, str(blk.get("text", "")), style)
        elif bt == "subheading":
            p = doc.add_paragraph(); _apply(p, prof.subheading)
            _run(p, str(blk.get("text", "")), prof.subheading)
        elif bt == "caption":
            p = doc.add_paragraph(); _apply(p, prof.caption)
            _run(p, str(blk.get("text", "")), prof.caption)
        elif bt in ("bullets", "numbered"):
            word_style = "List Number" if bt == "numbered" else "List Bullet"
            for item in blk.get("items", []):
                meta = item if isinstance(item, dict) else {"text": str(item)}
                lvl = int(meta.get("level", 1))
                sty = word_style if lvl <= 1 else f"{word_style} {min(lvl, 3)}"
                try:
                    p = doc.add_paragraph(style=sty)
                except KeyError:
                    p = doc.add_paragraph(style=word_style)
                base = prof.list_item
                _apply(p, base)
                _run(p, meta.get("text", ""), base,
                     bold=meta.get("bold", base.bold), italic=meta.get("italic", base.italic),
                     size=meta.get("size", base.size), color=meta.get("color", base.color))
        elif bt == "note":
            p = doc.add_paragraph(); _apply(p, prof.note)
            p.paragraph_format.left_indent = Inches(0.1)
            _shade(p, prof.note_bg)
            _run(p, "NOTE:  ", prof.note, bold=True)
            _run(p, str(blk.get("text", "")), prof.note)
        elif bt == "quote":
            p = doc.add_paragraph();
            qs = TextStyle(prof.paragraph.font, prof.paragraph.size, italic=True,
                           color="475569", left_indent=0.4, space_after=8)
            _apply(p, qs); _run(p, str(blk.get("text", "")), qs)
        elif bt == "table":
            _add_table(doc, prof, blk)
        else:  # paragraph
            p = doc.add_paragraph()
            base = prof.paragraph
            pstyle = TextStyle(
                base.font, blk.get("size", base.size), blk.get("bold", base.bold),
                blk.get("italic", base.italic), blk.get("color", base.color),
                blk.get("align", base.align), base.space_before,
                blk.get("space_after", base.space_after),
                float(blk.get("indent", base.left_indent)), base.line_spacing,
            )
            _apply(p, pstyle)
            runs = blk.get("runs")
            if runs:
                for rn in runs:
                    if not isinstance(rn, dict):
                        rn = {"text": str(rn)}
                    _run(p, rn.get("text", ""), pstyle,
                         bold=rn.get("bold", pstyle.bold), italic=rn.get("italic", pstyle.italic),
                         underline=rn.get("underline", False),
                         size=rn.get("size", pstyle.size), color=rn.get("color", pstyle.color))
            else:
                _run(p, str(blk.get("text", "")), pstyle)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _add_table(doc, prof: DocStyleProfile, blk: dict) -> None:
    headers = blk.get("headers", [])
    rows = blk.get("rows", [])
    ncol = len(headers) or (len(rows[0]) if rows else 1)
    table = doc.add_table(rows=0, cols=ncol)
    table.style = blk.get("word_style", "Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if headers:
        hr = table.add_row().cells
        for i, h in enumerate(headers):
            cell = hr[i]
            cell.paragraphs[0].text = ""
            _run(cell.paragraphs[0], str(h),
                 TextStyle(prof.heading2.font, 10, bold=True, color="FFFFFF"))
            _cell_fill(cell, prof.accent)
    for row in rows:
        rc = table.add_row().cells
        for i, val in enumerate(row[:ncol]):
            rc[i].paragraphs[0].text = ""
            _run(rc[i].paragraphs[0], str(val), TextStyle(prof.paragraph.font, 10))
    widths = blk.get("widths")
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths[:ncol]):
                row.cells[i].width = Inches(w)


def _cell_fill(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill.lstrip("#"))
    tc_pr.append(shd)


def available_profiles() -> list[str]:
    return list(PROFILES.keys())
