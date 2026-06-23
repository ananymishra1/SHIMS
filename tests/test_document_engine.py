"""Tests for the unified document engine (branded_base, template_manager, coa_engine, preview_engine)."""
from __future__ import annotations

import json
import tempfile
from pathlib import Path

import pytest

# Ensure DB tables exist before tests run
from shared.document_engine.template_manager import _init_tables, DocumentTemplateManager
from shared.document_engine.branded_base import BrandedPDF, DocumentLine, DocumentSection, FormatConfig
from shared.document_engine.coa_engine import COAEngine
from shared.document_engine.preview_engine import PreviewEngine

_init_tables()


class TestBrandedPDF:
    def test_build_basic_pdf(self):
        pdf = BrandedPDF(title="Test Document", doc_id="T-001")
        pdf.add_section(DocumentSection(
            title="Section A",
            lines=[
                DocumentLine(key="name", label="Name", value="Test Product"),
                DocumentLine(key="batch", label="Batch", value="B-001"),
            ]
        ))
        with tempfile.TemporaryDirectory() as td:
            path = pdf.build(Path(td) / "test.pdf")
            assert path.exists()
            assert path.stat().st_size > 1000

    def test_draft_watermark(self):
        config = FormatConfig(watermark_text="DRAFT", watermark_color="#FECACA")
        pdf = BrandedPDF(title="Draft", doc_id="D-001", format_config=config)
        pdf.add_section(DocumentSection(
            title="Test",
            lines=[DocumentLine(key="x", label="X", value="1")]
        ))
        with tempfile.TemporaryDirectory() as td:
            path = pdf.build(Path(td) / "draft.pdf")
            assert path.exists()

    def test_meta_table(self):
        pdf = BrandedPDF(title="Meta Test", doc_id="M-001")
        pdf.add_meta("Product", "Paracetamol")
        pdf.add_meta("Batch", "B-2026-001")
        with tempfile.TemporaryDirectory() as td:
            path = pdf.build(Path(td) / "meta.pdf")
            assert path.exists()

    def test_to_bytes(self):
        pdf = BrandedPDF(title="Bytes Test", doc_id="B-001")
        pdf.add_section(DocumentSection(
            title="Sec",
            lines=[DocumentLine(key="k", label="K", value="V")]
        ))
        data = pdf.to_bytes()
        assert isinstance(data, bytes)
        assert data.startswith(b"%PDF")


class TestDocumentTemplateManager:
    def test_create_and_get(self):
        tid = DocumentTemplateManager.create_template(name="Unit Test Template", kind="test")
        assert isinstance(tid, str) and tid.startswith("tmpl_")
        tpl = DocumentTemplateManager.get_template(tid)
        assert tpl is not None
        assert tpl["name"] == "Unit Test Template"
        assert tpl["kind"] == "test"
        DocumentTemplateManager.delete_template(tid)

    def test_list_by_kind(self):
        tid = DocumentTemplateManager.create_template(name="COA Test", kind="coa")
        coas = DocumentTemplateManager.list_templates(kind="coa")
        assert any(t["id"] == tid for t in coas)
        DocumentTemplateManager.delete_template(tid)

    def test_line_crud(self):
        tid = DocumentTemplateManager.create_template(name="Line Test", kind="test")
        lid = DocumentTemplateManager.add_line(tid, DocumentLine(key="pH", label="pH", type="number", required=True))
        assert lid.startswith("line_")

        tpl = DocumentTemplateManager.get_template(tid)
        assert len(tpl["lines"]) == 1
        assert tpl["lines"][0]["key"] == "pH"

        DocumentTemplateManager.update_line(lid, label="pH Value", spec="6.0-7.0")
        tpl = DocumentTemplateManager.get_template(tid)
        assert tpl["lines"][0]["label"] == "pH Value"
        assert tpl["lines"][0]["spec"] == "6.0-7.0"

        DocumentTemplateManager.delete_line(lid)
        tpl = DocumentTemplateManager.get_template(tid)
        assert len(tpl["lines"]) == 0
        DocumentTemplateManager.delete_template(tid)

    def test_move_line(self):
        tid = DocumentTemplateManager.create_template(name="Move Test", kind="test")
        l1 = DocumentTemplateManager.add_line(tid, DocumentLine(key="a", label="A", order=10))
        l2 = DocumentTemplateManager.add_line(tid, DocumentLine(key="b", label="B", order=20))
        l3 = DocumentTemplateManager.add_line(tid, DocumentLine(key="c", label="C", order=30))

        ok = DocumentTemplateManager.move_line(l3, 15)
        assert ok
        sections = DocumentTemplateManager.get_sections(tid)
        keys = [ln.key for ln in sections[0].lines]
        assert keys == ["a", "c", "b"]
        DocumentTemplateManager.delete_template(tid)

    def test_clone(self):
        tid = DocumentTemplateManager.create_template(name="Original", kind="test")
        DocumentTemplateManager.add_line(tid, DocumentLine(key="x", label="X"))
        new_id = DocumentTemplateManager.clone_template(tid, "Copy")
        assert new_id != tid
        copy = DocumentTemplateManager.get_template(new_id)
        assert copy["name"] == "Copy"
        assert len(copy["lines"]) == 1
        DocumentTemplateManager.delete_template(tid)
        DocumentTemplateManager.delete_template(new_id)

    def test_snapshot_restore(self):
        tid = DocumentTemplateManager.create_template(name="Snap", kind="test")
        DocumentTemplateManager.add_line(tid, DocumentLine(key="a", label="A"))
        sid = DocumentTemplateManager.snapshot(tid)

        # Modify
        DocumentTemplateManager.add_line(tid, DocumentLine(key="b", label="B"))
        tpl = DocumentTemplateManager.get_template(tid)
        assert len(tpl["lines"]) == 2

        # Restore
        DocumentTemplateManager.restore(sid)
        tpl = DocumentTemplateManager.get_template(tid)
        assert len(tpl["lines"]) == 1
        DocumentTemplateManager.delete_template(tid)


class TestCOAEngine:
    def test_generate_coa(self):
        with tempfile.TemporaryDirectory() as td:
            path = COAEngine.generate(
                product_name="Fluconazole IP",
                batch_number="B-2026-001",
                manufacturer="J.K. Lifecare",
                test_results=[
                    {"test": "Assay", "result": "99.2", "specification": "98.0-102.0", "method": "HPLC", "unit": "%"},
                    {"test": "pH", "result": "6.8", "specification": "6.0-7.0", "method": "USP", "unit": ""},
                ],
                output_path=Path(td) / "coa.pdf",
            )
            assert path.exists()
            assert path.stat().st_size > 1000

    def test_generate_from_form_data(self):
        with tempfile.TemporaryDirectory() as td:
            path = COAEngine.generate_from_form_data(
                {
                    "product_name": "Paracetamol",
                    "batch_number": "B-2026-002",
                    "assay": "99.5",
                    "ph": "6.5",
                },
                output_path=Path(td) / "coa_form.pdf",
            )
            assert path.exists()

    def test_preview_vs_finalize(self):
        form_data = {
            "product_name": "Test",
            "batch_number": "B-001",
            "description": "White powder",
        }
        with tempfile.TemporaryDirectory() as td:
            draft = COAEngine.preview(form_data, output_path=Path(td) / "draft.pdf")
            final = COAEngine.finalize(form_data, output_path=Path(td) / "final.pdf")
            assert draft.exists()
            assert final.exists()


class TestPreviewEngine:
    def test_preview_and_finalize(self):
        tid = DocumentTemplateManager.create_template(name="Preview Test", kind="business")
        DocumentTemplateManager.add_line(tid, DocumentLine(key="date", label="Date", type="date"))
        DocumentTemplateManager.add_line(tid, DocumentLine(key="recipient", label="To", type="text"))

        with tempfile.TemporaryDirectory() as td:
            preview_path = PreviewEngine.preview_document(
                tid, {"date": "2026-05-29", "recipient": "ABC Pharma"},
                output_path=Path(td) / "preview.pdf"
            )
            assert preview_path.exists()

            final_path = PreviewEngine.finalize_document(
                tid, {"date": "2026-05-29", "recipient": "ABC Pharma"},
                output_path=Path(td) / "final.pdf"
            )
            assert final_path.exists()

        DocumentTemplateManager.delete_template(tid)


class TestIntegration:
    def test_seed_defaults_creates_coa_template(self):
        DocumentTemplateManager.seed_defaults()
        coas = DocumentTemplateManager.list_templates(kind="coa")
        assert len(coas) >= 1
        # Verify it has lines
        tpl = DocumentTemplateManager.get_template(coas[0]["id"])
        assert len(tpl["lines"]) >= 10


class TestDocxExtraction:
    def test_extracts_paragraphs_and_tables(self):
        from docx import Document
        from backend.app.main import _extract_text_from_file

        with tempfile.TemporaryDirectory() as td:
            path = Path(td) / "sample.docx"
            doc = Document()
            doc.add_paragraph("Paragraph text.")
            table = doc.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = "Cell A"
            table.rows[0].cells[1].text = "Cell B"
            doc.save(path)

            text, error = _extract_text_from_file(path)
            assert error is None
            assert "Paragraph text." in text
            assert "Cell A | Cell B" in text

    def test_table_only_docx_returns_text(self):
        from docx import Document
        from backend.app.main import _extract_text_from_file

        with tempfile.TemporaryDirectory() as td:
            path = Path(td) / "table_only.docx"
            doc = Document()
            table = doc.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = "Only"
            table.rows[0].cells[1].text = "Table"
            doc.save(path)

            text, error = _extract_text_from_file(path)
            assert error is None
            assert "Only | Table" in text
