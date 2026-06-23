"""Regression coverage for the launch-hardening pass.

Locks in: the regulatory COA renderer, the rich-docx engine, Gmail send/reply
gating, and the Omni chemistry endpoints — the capabilities added/fixed for the
v11 production launch.
"""
from __future__ import annotations

from pathlib import Path

from fastapi.testclient import TestClient

from backend.app.main import app

client = TestClient(app)

FLU_SMILES = "OC(Cn1cncn1)(Cn1cncn1)c1ccc(F)cc1F"


# ── Regulatory COA renderer ─────────────────────────────────────────────────
def test_regulatory_coa_renders_pdf(tmp_path: Path):
    from shared.document_engine import coa_from_fields
    from shared.enterprise_pharma_core import FLUCONAZOLE_COA_FIELDS

    out = tmp_path / "coa.pdf"
    path = coa_from_fields(
        "Fluconazole IP", FLUCONAZOLE_COA_FIELDS,
        {"description": "White Crystalline Powder", "assay": "99.64%"},
        chemical_name="2-(2,4-difluorophenyl)-1,3-bis(1H-1,2,4-triazol-1-yl)propan-2-ol",
        batch_meta={"batch_no": "JKPD/01/26/001", "ar_no": "FP/26/001"},
        output_path=out,
    )
    data = Path(path).read_bytes()
    assert data.startswith(b"%PDF")
    assert len(data) > 2000


# ── Rich DOCX engine ────────────────────────────────────────────────────────
def test_rich_docx_styles_and_lists(tmp_path: Path):
    from shared.document_engine import build_docx, available_profiles

    assert {"regulatory", "corporate", "modern", "minimal"} <= set(available_profiles())
    out = tmp_path / "doc.docx"
    blocks = [
        {"type": "heading", "level": 1, "text": "Section"},
        {"type": "bullets", "items": ["a", {"text": "b", "bold": True}]},
        {"type": "numbered", "items": ["one", "two"]},
        {"type": "note", "text": "callout"},
        {"type": "table", "headers": ["k", "v"], "rows": [["x", "1"]]},
        {"type": "paragraph", "runs": [{"text": "bold", "bold": True}, " plain"]},
    ]
    build_docx("Title", blocks, profile="regulatory", output_path=out)
    from docx import Document
    doc = Document(str(out))
    styles = {p.style.name for p in doc.paragraphs}
    assert "List Bullet" in styles and "List Number" in styles
    assert len(doc.tables) == 1


# ── Gmail send/reply gating (no creds configured in test env) ───────────────
def test_gmail_send_requires_scope():
    r = client.post("/mailbox/gmail/send", json={"to": "x@example.com", "subject": "Hi", "body": "yo"})
    assert r.status_code == 428
    assert r.json()["status"] in {"scope_required", "needs_oauth"}


def test_gmail_reply_missing_message():
    r = client.post("/mailbox/gmail/reply", json={"message_id": "does-not-exist", "body": "hi"})
    assert r.status_code in (404, 428)


def test_mailbox_status_reports_send_capability():
    g = client.get("/mailbox/status").json()["gmail"]
    assert "send_enabled" in g and "client_secret_configured" in g


# ── Omni chemistry endpoints ────────────────────────────────────────────────
def test_chem_tools_list():
    r = client.get("/chem/tools")
    assert r.status_code == 200
    assert r.json()["ok"] and len(r.json()["tools"]) >= 1


def test_chem_verify_canonicalizes():
    r = client.post("/chem/verify", json={"smiles": FLU_SMILES})
    assert r.status_code == 200
    body = r.json()
    assert body["smiles"]["ok"] is True
    assert "hazards" in body


def test_chem_verify_rejects_empty():
    assert client.post("/chem/verify", json={"smiles": ""}).status_code == 400


def test_chem_ich_limits():
    r = client.post("/chem/ich", json={"impurity_pct": 0.13, "max_daily_dose_g": 0.4})
    assert r.status_code == 200 and r.json()["ok"] is True


def test_chem_config_resolves_ollama():
    # The integrated chem config must not crash resolving the local provider.
    from shared.shims_chem.config import get_config
    cfg = get_config()
    assert cfg.fast.url and cfg.fast.url.endswith("/v1")


# ── Rich-DOCX live endpoint ─────────────────────────────────────────────────
def test_rich_docx_endpoint_generates_and_serves():
    blocks = [
        {"type": "heading", "level": 1, "text": "Overview"},
        {"type": "numbered", "items": ["one", "two"]},
        {"type": "note", "text": "review required"},
    ]
    r = client.post("/documents/rich-docx", json={
        "title": "Endpoint SOP", "blocks": blocks, "profile": "regulatory", "subtitle": "SOP-1"})
    assert r.status_code == 200
    url = r.json()["url"]
    assert url.endswith(".docx")
    g = client.get(url)
    assert g.status_code == 200
    assert "wordprocessingml" in g.headers.get("content-type", "")


def test_rich_docx_endpoint_validates():
    assert client.post("/documents/rich-docx", json={"title": "x", "blocks": []}).status_code == 400
    assert client.post("/documents/rich-docx", json={
        "title": "x", "blocks": [{"type": "paragraph", "text": "hi"}], "profile": "nope"}).status_code == 400


def test_document_profiles_endpoint():
    profs = client.get("/documents/profiles").json()["profiles"]
    assert {"regulatory", "corporate", "modern", "minimal"} <= set(profs)


# ── RDKit is active (chem accuracy upgrade) ─────────────────────────────────
def test_rdkit_active_and_rejects_empty():
    from shared.shims_chem.verifier import smiles as sm
    assert sm._RDKIT is True
    assert sm.validate_smiles("").ok is False  # empty must never validate


# ── Lifespan migration (no deprecated on_event) ─────────────────────────────
def test_no_on_event_deprecation():
    import warnings
    import importlib
    with warnings.catch_warnings(record=True) as caught:
        warnings.simplefilter("always")
        importlib.reload(importlib.import_module("backend.app.main"))
    assert not [w for w in caught if "on_event" in str(w.message)]


# ── STT model switch (Settings UI backend) ──────────────────────────────────
def test_stt_models_listing_and_switch():
    listing = client.get("/stt/models").json()
    assert listing["ok"] and "active" in listing and isinstance(listing["models"], list)
    original = listing["active"]
    try:
        # Switching to a built-in size name is always accepted.
        r = client.post("/stt/model", json={"model": "base"})
        assert r.status_code == 200 and r.json()["active"] == "base"
        assert client.get("/stt/health").json()["model"] == "base"
        # Unknown model is rejected.
        assert client.post("/stt/model", json={"model": "bogus-xyz"}).status_code == 400
    finally:
        # Restore the prior selection so the test never disturbs runtime config.
        client.post("/stt/model", json={"model": original})


# ── Live Ollama/LLM smoke test (skips cleanly when offline) ─────────────────
def test_ollama_health_smoke():
    import pytest
    health = client.get("/health").json()
    assert "ollama_host" in health
    if not health.get("ollama_online"):
        pytest.skip("Ollama not reachable in this environment")
    # When online, the model list should be populated.
    assert isinstance(health.get("models"), list) and health["models"]
